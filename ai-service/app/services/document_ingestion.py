import base64
from collections import Counter
from io import BytesIO
import logging
from pathlib import Path
import re
import shutil
import subprocess
import tempfile
import zipfile
from xml.etree import ElementTree

import httpx

from app.config import settings
from app.models.schemas import ContractIngestionRequest, DocumentExtractionMetadata, EvaluationMetricSet, StructuredDocument
from app.services.document_preprocessing_engine import document_preprocessing_engine
from app.services.document_understanding_engine import document_understanding_engine
from app.services.file_router import file_router
from app.services.ocr_pipeline import ocr_pipeline
from app.services.pdf_converter import pdf_converter


PAGE_NUMBER_PATTERN = re.compile(r"^\d+\s*$")
LEGAL_MARKER_PATTERN = re.compile(
    r"\b(section|clause|article|chapter|party|term|termination|liability|payment|indemn|confidential|governing law)\b",
    re.IGNORECASE,
)
BROKEN_WORD_PATTERN = re.compile(r"(\w)-\s*\n\s*(\w)")
MULTISPACE_PATTERN = re.compile(r"[ \t]{2,}")
NOISE_PATTERN = re.compile(r"^[\W_]+$")
logger = logging.getLogger("contractsense.ingestion")


class DocumentIngestionService:
    def __init__(self) -> None:
        self._latest_ocr_blocks = []

    def extract(self, payload: ContractIngestionRequest) -> tuple[str, DocumentExtractionMetadata]:
        text, metadata, _, _ = self.extract_with_structure(payload)
        return text, metadata

    def extract_with_structure(
        self,
        payload: ContractIngestionRequest,
    ) -> tuple[str, DocumentExtractionMetadata, StructuredDocument | None, EvaluationMetricSet | None]:
        self._latest_ocr_blocks = []
        if payload.contractText.strip():
            text = self._normalize_legal_text(payload.contractText.strip())
            logger.info("RAW OCR OUTPUT:\n%s", text[:4000])
            metadata = self._metadata(
                payload,
                file_router.detect_kind(payload.fileName, payload.mediaType),
                "direct_text_payload",
                0.95,
                text,
                ["Used provided contract text directly.", "Applied legal-text normalization."],
            )
            structured, evaluation = document_understanding_engine.process_text(text, metadata.documentKind, metadata.extractionMethod)
            metadata.structuredPages = len(structured.pages)
            metadata.layoutElements = sum(len(page.elements) for page in structured.pages)
            metadata.structureConfidence = max(metadata.structureConfidence, evaluation.documentStructureAccuracy)
            return text, metadata, structured, evaluation

        raw_bytes = base64.b64decode(payload.documentBase64) if payload.documentBase64 else b""
        document_kind = file_router.detect_kind(payload.fileName, payload.mediaType)
        if document_kind == "txt":
            text = self._normalize_legal_text(self._decode_text(raw_bytes))
            logger.info("RAW OCR OUTPUT:\n%s", text[:4000])
            metadata = self._metadata(payload, document_kind, "text_decode", self._text_quality_score(text), text)
            structured, evaluation = document_understanding_engine.process_text(text, metadata.documentKind, metadata.extractionMethod)
            metadata.structuredPages = len(structured.pages)
            metadata.layoutElements = sum(len(page.elements) for page in structured.pages)
            metadata.structureConfidence = max(metadata.structureConfidence, evaluation.documentStructureAccuracy)
            return text, metadata, structured, evaluation
        if document_kind == "docx":
            text, metadata = self._extract_docx(raw_bytes, payload)
            return self._augment_with_structure(raw_bytes, payload, text, metadata)
        if document_kind == "doc":
            text, metadata = self._extract_doc(raw_bytes, payload)
            return self._augment_with_structure(raw_bytes, payload, text, metadata)
        if document_kind == "pdf":
            text, metadata = self._extract_pdf(raw_bytes, payload)
            return self._augment_with_structure(raw_bytes, payload, text, metadata)
        if document_kind == "image":
            text, metadata = self._extract_image(raw_bytes, payload)
            return self._augment_with_structure(raw_bytes, payload, text, metadata)
        text = self._normalize_legal_text(self._decode_text(raw_bytes))
        return text, self._metadata(
            payload,
            document_kind,
            "best_effort_binary_decode",
            max(0.2, self._text_quality_score(text) * 0.5),
            text,
            ["Unknown format handled with lossy decoding."],
        ), None, None

    def _extract_docx(self, raw_bytes: bytes, payload: ContractIngestionRequest) -> tuple[str, DocumentExtractionMetadata]:
        notes: list[str] = []
        try:
            from docx import Document  # type: ignore[import-not-found]

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                temp_file.write(raw_bytes)
                temp_path = temp_file.name
            document = Document(temp_path)
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            text = self._normalize_legal_text("\n".join(paragraphs))
            logger.info("RAW OCR OUTPUT:\n%s", text[:4000] if text else "<empty>")
            return text, self._metadata(payload, "docx", "python_docx_parse", max(0.84, self._text_quality_score(text)), text, notes)
        except Exception:
            try:
                with zipfile.ZipFile(BytesIO(raw_bytes)) as archive:
                    xml_bytes = archive.read("word/document.xml")
                root = ElementTree.fromstring(xml_bytes)
                paragraphs: list[str] = []
                current = []
                for node in root.iter():
                    tag = node.tag.lower()
                    if tag.endswith("}t") and node.text:
                        current.append(node.text)
                    elif tag.endswith("}p"):
                        paragraph = "".join(current).strip()
                        if paragraph:
                            paragraphs.append(paragraph)
                        current = []
                text = self._normalize_legal_text("\n".join(paragraphs))
                logger.info("RAW OCR OUTPUT:\n%s", text[:4000] if text else "<empty>")
                notes.append("DOCX parsing fell back to XML extraction.")
                return text, self._metadata(payload, "docx", "docx_xml_parse", max(0.72, self._text_quality_score(text)), text, notes)
            except Exception:
                notes.append("DOCX parsing fell back to lossy text decoding.")
                text = self._normalize_legal_text(self._decode_text(raw_bytes))
                logger.info("RAW OCR OUTPUT:\n%s", text[:4000] if text else "<empty>")
                return text, self._metadata(payload, "docx", "docx_fallback_decode", 0.42, text, notes)

    def _extract_doc(self, raw_bytes: bytes, payload: ContractIngestionRequest) -> tuple[str, DocumentExtractionMetadata]:
        notes: list[str] = []
        with tempfile.TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / payload.fileName
            source_path.write_bytes(raw_bytes)

            antiword = shutil.which("antiword")
            if antiword:
                try:
                    completed = subprocess.run(
                        [antiword, str(source_path)],
                        capture_output=True,
                        text=True,
                        check=False,
                        timeout=30,
                    )
                    text = self._normalize_legal_text(completed.stdout.strip())
                    if text:
                        notes.append("Legacy DOC parsed with antiword.")
                        return text, self._metadata(payload, "doc", "antiword_doc_parse", max(0.72, self._text_quality_score(text)), text, notes)
                except Exception:
                    notes.append("antiword parsing failed.")

            soffice = shutil.which("soffice")
            if soffice:
                try:
                    subprocess.run(
                        [soffice, "--headless", "--convert-to", "txt:Text", "--outdir", temp_dir, str(source_path)],
                        capture_output=True,
                        text=True,
                        check=False,
                        timeout=60,
                    )
                    converted_path = source_path.with_suffix(".txt")
                    if converted_path.exists():
                        text = self._normalize_legal_text(converted_path.read_text(encoding="utf-8", errors="ignore").strip())
                        if text:
                            notes.append("Legacy DOC converted through LibreOffice.")
                            return text, self._metadata(payload, "doc", "libreoffice_doc_convert", max(0.76, self._text_quality_score(text)), text, notes)
                except Exception:
                    notes.append("LibreOffice DOC conversion failed.")

        notes.append("No DOC parser was available; used lossy text decoding.")
        text = self._normalize_legal_text(self._decode_text(raw_bytes))
        return text, self._metadata(payload, "doc", "doc_fallback_decode", 0.22, text, notes)

    def _extract_pdf(self, raw_bytes: bytes, payload: ContractIngestionRequest) -> tuple[str, DocumentExtractionMetadata]:
        notes: list[str] = []
        candidates: list[tuple[str, str, list[str]]] = []

        if file_router.pdf_has_text_layer(raw_bytes):
            gemini_text = self._extract_with_gemini(
                raw_bytes,
                "application/pdf",
                (
                    "Transcribe this legal PDF in reading order. Preserve clause numbering, headings, section boundaries, "
                    "and paragraph separation. Return only the extracted document text."
                ),
            )
            if gemini_text:
                candidates.append((self._normalize_legal_text(gemini_text), "gemini_pdf_understanding", ["Gemini document understanding extracted the PDF content."]))

            try:
                from pypdf import PdfReader  # type: ignore[import-not-found]

                reader = PdfReader(BytesIO(raw_bytes))
                pypdf_text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
                if pypdf_text:
                    candidates.append((self._normalize_legal_text(pypdf_text), "pypdf_text_extraction", ["Native PDF text layer extracted with pypdf."]))
            except Exception:
                notes.append("pypdf extraction failed.")

            pymupdf_blocks = self._extract_pdf_blocks_with_pymupdf(raw_bytes, payload.fileName)
            if pymupdf_blocks:
                candidates.append((self._normalize_legal_text(pymupdf_blocks), "pymupdf_block_extraction", ["PyMuPDF block-based layout extraction succeeded."]))

            pymupdf_text = self._extract_pdf_text_with_pymupdf(raw_bytes, payload.fileName)
            if pymupdf_text:
                candidates.append((self._normalize_legal_text(pymupdf_text), "pymupdf_text_extraction", ["PyMuPDF plain-text extraction succeeded."]))

        if candidates:
            best_text, best_method, best_notes, best_score = self._pick_best_candidate(candidates)
            if best_text.strip():
                logger.info("RAW OCR OUTPUT:\n%s", best_text[:4000])
                confidence = max(0.45, min(0.96, best_score))
                return best_text, self._metadata(payload, "pdf", best_method, confidence, best_text, best_notes + notes)

        pages = pdf_converter.convert_to_images(raw_bytes, payload.fileName)
        text, blocks, ocr_notes = ocr_pipeline.run_pages(pages)
        self._latest_ocr_blocks = blocks
        if text.strip():
            normalized = self._normalize_legal_text(text)
            logger.info("RAW OCR OUTPUT:\n%s", normalized[:4000])
            return normalized, self._metadata(payload, "pdf", "ocr_scanned_pdf", max(0.62, self._text_quality_score(normalized)), normalized, notes + ocr_notes + ["Scanned PDF pages were converted to images and processed with OCR."])

        logger.warning("RAW OCR OUTPUT: empty after PDF OCR fallback.")
        return "", self._metadata(payload, "pdf", "pdf_ocr_failed", 0.0, "", notes + ocr_notes + ["No PDF extractor or OCR engine produced usable text."])

    def _extract_image(self, raw_bytes: bytes, payload: ContractIngestionRequest) -> tuple[str, DocumentExtractionMetadata]:
        notes: list[str] = []
        pages = document_preprocessing_engine.preprocess(raw_bytes, payload.fileName, "image").pages
        text, blocks, ocr_notes = ocr_pipeline.run_pages(pages)
        self._latest_ocr_blocks = blocks
        if text.strip():
            normalized = self._normalize_legal_text(text)
            logger.info("RAW OCR OUTPUT:\n%s", normalized[:4000])
            return normalized, self._metadata(payload, "image", "paddleocr_image_pipeline", max(0.66, self._text_quality_score(normalized)), normalized, notes + ocr_notes)

        notes.append("Image OCR returned no text after primary and fallback attempts.")
        logger.warning("RAW OCR OUTPUT: empty after image OCR pipeline.")
        return "", self._metadata(payload, "image", "image_ocr_unavailable", 0.05, "", notes)

    def _augment_with_structure(
        self,
        raw_bytes: bytes,
        payload: ContractIngestionRequest,
        text: str,
        metadata: DocumentExtractionMetadata,
    ) -> tuple[str, DocumentExtractionMetadata, StructuredDocument | None, EvaluationMetricSet | None]:
        try:
            if self._latest_ocr_blocks:
                structured, evaluation = document_understanding_engine.process_ocr(
                    blocks=self._latest_ocr_blocks,
                    text=text,
                    document_kind=metadata.documentKind,
                    extraction_method=metadata.extractionMethod,
                )
            else:
                structured, evaluation = document_understanding_engine.process(
                    raw_bytes=raw_bytes,
                    file_name=payload.fileName,
                    document_kind=metadata.documentKind,
                    extraction_method=metadata.extractionMethod,
                    normalized_text=text,
                )
        except Exception as error:
            metadata.notes.append(f"Document understanding engine fallback activated: {error}")
            return text, metadata, None, None

        metadata.structuredPages = len(structured.pages)
        metadata.layoutElements = sum(len(page.elements) for page in structured.pages)
        metadata.tablesDetected = len(structured.tables)
        metadata.signaturesDetected = len(structured.signatures)
        metadata.annotationsDetected = len(structured.annotations)
        metadata.structureConfidence = max(metadata.structureConfidence, evaluation.documentStructureAccuracy)
        if structured.rawText:
            text = structured.rawText
            metadata.extractedCharacters = len(text)
            metadata.textPreview = text[:400]
        if not text.strip():
            logger.warning("RAW OCR OUTPUT: still empty before downstream analysis.")
        metadata.notes.extend(structured.reconstructionNotes[:4])
        return text, metadata, structured, evaluation

    def _document_kind(self, file_name: str, media_type: str) -> str:
        lowered_name = file_name.lower()
        lowered_media = media_type.lower()
        if lowered_name.endswith(".pdf") or "pdf" in lowered_media:
            return "pdf"
        if lowered_name.endswith(".docx") or "wordprocessingml" in lowered_media:
            return "docx"
        if lowered_name.endswith(".doc"):
            return "doc"
        if lowered_name.endswith((".jpg", ".jpeg", ".png")) or lowered_media.startswith("image/"):
            return "image"
        if lowered_name.endswith(".txt") or lowered_media.startswith("text/"):
            return "txt"
        return "unknown"

    def _decode_text(self, raw_bytes: bytes) -> str:
        if not raw_bytes:
            return ""
        for encoding in ("utf-8", "utf-16", "latin-1"):
            try:
                return raw_bytes.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
        return raw_bytes.decode("utf-8", errors="ignore").strip()

    def _metadata(
        self,
        payload: ContractIngestionRequest,
        document_kind: str,
        extraction_method: str,
        confidence: float,
        text: str,
        notes: list[str] | None = None,
    ) -> DocumentExtractionMetadata:
        return DocumentExtractionMetadata(
            documentKind=document_kind,  # type: ignore[arg-type]
            extractionMethod=extraction_method,
            sourceFilename=payload.fileName,
            mediaType=payload.mediaType,
            structureConfidence=max(0.0, min(1.0, confidence)),
            extractedCharacters=len(text),
            textPreview=text[:400],
            notes=notes or [],
            structuredPages=0,
            layoutElements=0,
            tablesDetected=0,
            signaturesDetected=0,
            annotationsDetected=0,
        )

    def _image_suffix(self, file_name: str) -> str:
        if "." not in file_name:
            return ".png"
        return "." + file_name.rsplit(".", 1)[1]

    def _extract_pdf_text_with_pymupdf(self, raw_bytes: bytes, file_name: str) -> str:
        try:
            import fitz  # type: ignore[import-not-found]
        except Exception:
            return ""

        with tempfile.TemporaryDirectory() as temp_dir:
            pdf_path = Path(temp_dir) / file_name
            pdf_path.write_bytes(raw_bytes)
            document = fitz.open(str(pdf_path))
            page_texts = []
            for page_index in range(len(document)):
                page = document.load_page(page_index)
                page_texts.append((page.get_text("text") or "").strip())
            document.close()
            return "\n\n".join(text for text in page_texts if text).strip()

    def _extract_pdf_blocks_with_pymupdf(self, raw_bytes: bytes, file_name: str) -> str:
        try:
            import fitz  # type: ignore[import-not-found]
        except Exception:
            return ""

        with tempfile.TemporaryDirectory() as temp_dir:
            pdf_path = Path(temp_dir) / file_name
            pdf_path.write_bytes(raw_bytes)
            document = fitz.open(str(pdf_path))
            pages: list[str] = []
            for page_index in range(len(document)):
                page = document.load_page(page_index)
                block_texts = self._ordered_page_blocks(page)
                if block_texts:
                    pages.append("\n".join(block_texts))
            document.close()
            return "\n\n".join(pages).strip()

    def _pick_best_candidate(self, candidates: list[tuple[str, str, list[str]]]) -> tuple[str, str, list[str], float]:
        scored: list[tuple[float, str, str, list[str]]] = []
        for text, method, notes in candidates:
            score = self._text_quality_score(text)
            scored.append((score, text, method, notes))
        scored.sort(key=lambda item: item[0], reverse=True)
        best_score, best_text, best_method, best_notes = scored[0]
        return best_text, best_method, best_notes, best_score

    def _normalize_legal_text(self, text: str) -> str:
        if not text:
            return ""
        normalized = text.replace("\r\n", "\n").replace("\r", "\n")
        normalized = BROKEN_WORD_PATTERN.sub(r"\1\2", normalized)
        lines = [MULTISPACE_PATTERN.sub(" ", line).strip() for line in normalized.split("\n")]
        lines = self._remove_repeated_headers(lines)
        lines = self._repair_fragmented_lines(lines)
        rebuilt: list[str] = []
        buffer = ""

        for raw_line in lines:
            line = raw_line.strip()
            if not line or NOISE_PATTERN.match(line) or PAGE_NUMBER_PATTERN.match(line):
                continue
            if self._is_heading_line(line) or self._is_clause_start(line):
                if buffer:
                    rebuilt.append(buffer.strip())
                    buffer = ""
                rebuilt.append(line)
                continue
            if not buffer:
                buffer = line
                continue
            if self._should_join(buffer, line):
                buffer = f"{buffer} {line}".strip()
            else:
                rebuilt.append(buffer.strip())
                buffer = line

        if buffer:
            rebuilt.append(buffer.strip())

        compact: list[str] = []
        for line in rebuilt:
            if compact and line == compact[-1]:
                continue
            compact.append(line)
        return "\n".join(compact).strip()

    def _remove_repeated_headers(self, lines: list[str]) -> list[str]:
        stripped = [line.strip() for line in lines if line.strip()]
        counts = Counter(stripped)
        repeated = {
            line
            for line, count in counts.items()
            if count >= 3 and len(line.split()) <= 10 and (line.isupper() or "page" in line.lower())
        }
        return [line for line in lines if line.strip() not in repeated]

    def _is_heading_line(self, line: str) -> bool:
        upper = line.isupper() and 1 <= len(line.split()) <= 14
        return upper or bool(re.match(r"^(chapter|part|article|schedule|appendix)\b", line, re.IGNORECASE))

    def _is_clause_start(self, line: str) -> bool:
        return bool(re.match(r"^\d+[A-Z]?(?:\.\d+)*[\).]?\s+\S+", line) or re.match(r"^[A-Z][\).]\s+\S+", line))

    def _should_join(self, previous: str, current: str) -> bool:
        if self._is_heading_line(current) or self._is_clause_start(current):
            return False
        if previous.endswith((".", ";", ":")):
            return False
        if len(current.split()) <= 3 and current[0].isupper():
            return False
        return True

    def _ordered_page_blocks(self, page) -> list[str]:
        try:
            blocks = page.get_text("dict").get("blocks", [])
        except Exception:
            return []
        text_blocks = [block for block in blocks if block.get("type") == 0]
        if not text_blocks:
            return []

        page_width = float(page.rect.width or 1.0)
        column_cutoff = page_width * 0.18
        left_blocks = [block for block in text_blocks if float(block.get("bbox", [0.0])[0]) < column_cutoff]
        right_blocks = [block for block in text_blocks if float(block.get("bbox", [0.0])[0]) >= column_cutoff]
        ordered_columns = [left_blocks, right_blocks] if left_blocks and right_blocks else [text_blocks]

        ordered_texts: list[str] = []
        for column_blocks in ordered_columns:
            for block in sorted(column_blocks, key=lambda item: (round(item.get("bbox", [0, 0])[1], 1), round(item.get("bbox", [0, 0])[0], 1))):
                lines: list[str] = []
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    line_text = "".join(str(span.get("text") or "") for span in spans).strip()
                    if line_text:
                        lines.append(line_text)
                paragraph = self._normalize_legal_text("\n".join(lines)).strip()
                if paragraph:
                    ordered_texts.append(paragraph)
        return ordered_texts

    def _repair_fragmented_lines(self, lines: list[str]) -> list[str]:
        repaired: list[str] = []
        for line in lines:
            if not line:
                continue
            if repaired and self._should_attach_to_previous(repaired[-1], line):
                repaired[-1] = f"{repaired[-1]} {line}".strip()
            else:
                repaired.append(line)
        return repaired

    def _should_attach_to_previous(self, previous: str, current: str) -> bool:
        if self._is_heading_line(current) or self._is_clause_start(current):
            return False
        if previous.endswith((",", "(", "-", "/")):
            return True
        if current[:1].islower():
            return True
        if len(current.split()) <= 2 and not current.endswith((".", ";", ":")):
            return True
        if previous and previous[-1].isalnum() and current and current[0].isalnum() and not previous.endswith((".", ";", ":")):
            return len(current.split()) <= 6
        return False

    def _text_quality_score(self, text: str) -> float:
        if not text:
            return 0.0
        alpha_chars = sum(1 for character in text if character.isalpha())
        printable_chars = sum(1 for character in text if character.isprintable())
        lines = [line for line in text.splitlines() if line.strip()]
        line_count = len(lines)
        unique_tokens = len({token.lower() for token in text.split() if len(token) > 2})
        alpha_ratio = alpha_chars / max(1, len(text))
        printable_ratio = printable_chars / max(1, len(text))
        lexical_bonus = min(0.18, unique_tokens / 450.0)
        structure_bonus = min(0.18, line_count / 180.0)
        legal_marker_bonus = min(0.18, len(LEGAL_MARKER_PATTERN.findall(text)) / 80.0)
        garbage_penalty = 0.0
        short_lines = len([line for line in lines if len(line.split()) <= 2])
        if line_count:
            garbage_penalty += min(0.12, short_lines / line_count * 0.18)
        if "cid:" in text.lower():
            garbage_penalty += 0.1
        return round(max(0.0, min(1.0, 0.34 * alpha_ratio + 0.22 * printable_ratio + lexical_bonus + structure_bonus + legal_marker_bonus - garbage_penalty)), 2)

    def _extract_with_gemini(self, raw_bytes: bytes, mime_type: str, prompt: str) -> str:
        if not settings.gemini_api_key or not raw_bytes:
            return ""
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{settings.gemini_model}:generateContent"
        payload = {
            "contents": [
                {
                    "parts": [
                        {
                            "inline_data": {
                                "mime_type": mime_type,
                                "data": base64.b64encode(raw_bytes).decode("utf-8"),
                            }
                        },
                        {"text": prompt},
                    ]
                }
            ],
            "generationConfig": {
                "temperature": 0.0,
                "topP": 0.9,
                "maxOutputTokens": 8192,
            },
        }
        try:
            response = httpx.post(url, params={"key": settings.gemini_api_key}, json=payload, timeout=120.0)
            response.raise_for_status()
            data = response.json()
            candidates = data.get("candidates", [])
            if not candidates:
                return ""
            parts = candidates[0].get("content", {}).get("parts", [])
            text = "\n".join(part.get("text", "").strip() for part in parts if part.get("text"))
            return text.strip()
        except Exception:
            return ""


document_ingestion_service = DocumentIngestionService()
